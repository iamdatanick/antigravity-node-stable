"""Document processing — extract text from uploads and chunk for RAG."""

import io
import logging
import os

logger = logging.getLogger("antigravity.document_processor")


def process_document(file_path: str) -> str:
    """Extract text from a file on disk (PDF, DOCX, or plain text)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        import pypdf

        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages])
    elif ext == ".docx":
        import docx

        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        with open(file_path, encoding="utf-8") as f:
            return f.read()


def process_bytes(data: bytes, filename: str) -> str:
    """Extract text from in-memory bytes. Used by the upload endpoint."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif ext == ".docx":
        import docx

        doc = docx.Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return data.decode("utf-8", errors="replace")


def chunk_text(text: str, size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    chunks = []
    for i in range(0, len(text), size - overlap):
        chunks.append(text[i : i + size])
    return chunks
